"""Document processing service."""
import hashlib
from typing import Optional
from docx import Document
from docx.opc.exceptions import PackageNotFoundError


class DocumentService:
    """Service for processing document files."""
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """
        Extract text content from Word document.
        
        Args:
            file_path: Path to .docx file
        
        Returns:
            Extracted text content
        
        Raises:
            ValueError: If file is not valid DOCX
        """
        try:
            doc = Document(file_path)
            
            # Extract paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Extract tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_texts.append(row_text)
            
            # Combine
            all_text = "\n\n".join(paragraphs)
            if table_texts:
                all_text += "\n\n--- Tables ---\n\n" + "\n".join(table_texts)
            
            return all_text
        
        except PackageNotFoundError as e:
            raise ValueError(f"Invalid DOCX file: {e}") from e
    
    @staticmethod
    def calculate_file_hash(file_path: str) -> str:
        """
        Calculate SHA-256 hash of file.
        
        Args:
            file_path: Path to file
        
        Returns:
            Hex-encoded SHA-256 hash
        """
        sha256_hash = hashlib.sha256()
        
        with open(file_path, "rb") as f:
            # Read in 64KB chunks
            for byte_block in iter(lambda: f.read(65536), b""):
                sha256_hash.update(byte_block)
        
        return sha256_hash.hexdigest()
    
    @staticmethod
    def parse_filename_metadata(filename: str) -> dict:
        """
        Parse metadata from filename using heuristics.
        
        Common patterns:
        - YYYY-MM-DD_Company_Type.ext
        - Meeting_with_Company_YYYY-MM-DD.ext
        - IK_Company_möte_YYYYMMDD_HH-MM__Duplicate_Info.ext
        
        Args:
            filename: File name
        
        Returns:
            Dict with parsed metadata including cleaned_name
        """
        import re
        from datetime import datetime
        
        metadata = {}
        
        # Try to extract date (YYYY-MM-DD or YYYYMMDD)
        date_match = re.search(r"(\d{4})-?(\d{2})-?(\d{2})", filename)
        if date_match:
            try:
                date_str = f"{date_match.group(1)}-{date_match.group(2)}-{date_match.group(3)}"
                metadata["date"] = date_str
            except ValueError:
                pass
        
        # Try to extract time (HH:MM or HH-MM)
        time_match = re.search(r"(\d{2})[-:](\d{2})", filename)
        if time_match:
            metadata["time"] = f"{time_match.group(1)}:{time_match.group(2)}"
        
        # Try to extract company name (simple heuristic: capitalized word)
        company_match = re.search(r"([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)", filename)
        if company_match:
            metadata["company"] = company_match.group(1)
        
        # Try to extract meeting type
        meeting_types = ["kickoff", "standup", "review", "planning", "retrospective", "demo", 
                         "möte", "meeting", "intro", "team meeting", "partner"]
        filename_lower = filename.lower()
        for mtype in meeting_types:
            if mtype in filename_lower:
                metadata["type"] = mtype
                break
        
        # Generate clean name
        metadata["cleaned_name"] = DocumentService.clean_document_name(filename)
        
        return metadata
    
    @staticmethod
    def clean_document_name(filename: str) -> str:
        """
        Clean and standardize document name for display.
        
        Examples:
        - "IK_Disruptive_Ventures_möte_20231005_10-05__IK, Disruptive Ventures"
          -> "Disruptive Ventures Meeting - Oct 5, 2023"
        
        - "Möte_2023-10-04_Serge_Guelnoji_Peo__2023-10-04__serge _ guelnoji _ peo"
          -> "Meeting with Serge Guelnoji - Oct 4, 2023"
        
        - "Online_Partner_disruptiveventures_Gemini Enterprise SKU, Online Part"
          -> "Gemini Enterprise SKU - Online Partner"
        
        Args:
            filename: Original filename
        
        Returns:
            Cleaned, human-readable name
        """
        import re
        from datetime import datetime
        
        # Remove file extension
        name_no_ext = re.sub(r'\.[^.]+$', '', filename)
        
        # Split by common delimiters
        parts = re.split(r'[_\-\s]+|__+', name_no_ext)
        
        # Remove empty parts and common prefixes
        filtered_parts = []
        seen_parts = set()
        common_prefixes = {'ik', 'möte', 'meeting', 'online', 'partner'}
        
        date_str = None
        time_str = None
        names = []
        meeting_type = None
        
        for part in parts:
            part_clean = part.strip().strip(',')
            part_lower = part_clean.lower()
            
            # Skip empty or very short parts
            if not part_clean or len(part_clean) < 2:
                continue
            
            # Extract date
            date_match = re.match(r'(\d{4})[-/]?(\d{2})[-/]?(\d{2})', part_clean)
            if date_match:
                try:
                    year, month, day = date_match.groups()
                    date_obj = datetime(int(year), int(month), int(day))
                    date_str = date_obj.strftime("%b %d, %Y")
                except ValueError:
                    pass
                continue
            
            # Extract time
            time_match = re.match(r'(\d{2})[-:](\d{2})', part_clean)
            if time_match and not date_str:  # Only if not part of date
                time_str = f"{time_match.group(1)}:{time_match.group(2)}"
                continue
            
            # Identify meeting types
            if part_lower in ['möte', 'meeting', 'intro', 'call', 'discussion']:
                meeting_type = 'Meeting'
                continue
            
            # Skip common prefixes and duplicates
            if part_lower in common_prefixes:
                continue
            
            # Check for duplicates (case-insensitive)
            if part_lower not in seen_parts:
                seen_parts.add(part_lower)
                
                # Capitalize properly
                if re.match(r'^[A-Z][a-z]+(?:[A-Z][a-z]+)*$', part_clean):
                    # Already properly capitalized (e.g., "DisruptiveVentures")
                    # Add spaces before capitals
                    part_clean = re.sub(r'([a-z])([A-Z])', r'\1 \2', part_clean)
                elif part_clean.isupper() or part_clean.islower():
                    # All caps or all lower -> Title case
                    part_clean = part_clean.title()
                
                names.append(part_clean)
        
        # Build final name
        result_parts = []
        
        # Add main subject (filter out generic terms from beginning)
        main_names = [n for n in names if n.lower() not in ['disruptive', 'ventures', 'disruptive ventures']]
        if not main_names:
            main_names = names
        
        if main_names:
            # Primary subject
            primary = main_names[0]
            result_parts.append(primary)
            
            # Add meeting type if identified
            if meeting_type:
                result_parts.append(meeting_type)
            
            # Add additional context if available and not redundant
            if len(main_names) > 1:
                additional = ', '.join(main_names[1:3])  # Limit to avoid clutter
                if additional.lower() not in primary.lower():
                    result_parts.append(f"({additional})")
        elif meeting_type:
            result_parts.append(meeting_type)
        
        # Build final string
        final_name = ' - '.join(result_parts) if result_parts else 'Document'
        
        # Add date if found
        if date_str:
            final_name += f" - {date_str}"
        
        # Add time if found and no date
        if time_str and not date_str:
            final_name += f" at {time_str}"
        
        # Clean up any remaining issues
        final_name = re.sub(r'\s+', ' ', final_name)  # Multiple spaces
        final_name = re.sub(r'\s*-\s*-\s*', ' - ', final_name)  # Multiple dashes
        final_name = final_name.strip(' -')
        
        # Limit length
        if len(final_name) > 100:
            final_name = final_name[:97] + '...'
        
        return final_name




