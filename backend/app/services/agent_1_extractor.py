"""
Agent 1: Extractor Service
Purpose: Extract structured data from documents with ZERO interpretation
Rule: NEVER infer, assume, or fill gaps - extract ONLY what is explicitly present
"""

import hashlib
import re
from typing import Optional, Dict, Any, List
from enum import Enum
from datetime import datetime
import io

from pydantic import BaseModel, Field
import pypdf2
import pdfplumber
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
import magic

from app.config import settings
from app.services.document import DocumentService


class EntityType(str, Enum):
    """Types of entities that can be extracted."""
    COMPANY = "company"
    PERSON = "person"
    MONEY = "money"
    DATE = "date"
    PERCENTAGE = "percentage"
    METRIC = "metric"
    URL = "url"
    EMAIL = "email"
    PHONE = "phone"


class ExtractedEntity(BaseModel):
    """An entity extracted from the document."""
    type: EntityType
    value: str
    source_location: str  # e.g., "page 3", "line 45", "cell B2"
    confidence: float = Field(ge=0.0, le=1.0)
    context: Optional[str] = None  # Surrounding text for verification


class ExtractedTable(BaseModel):
    """A table extracted from the document."""
    location: str  # "page 3" or "section 2.1"
    headers: List[str]
    rows: List[List[str]]
    confidence: float = Field(ge=0.0, le=1.0)
    malformed: bool = False


class Ambiguity(BaseModel):
    """A flagged ambiguous section."""
    location: str
    issue: str
    raw_text: Optional[str] = None


class ExtractionResult(BaseModel):
    """Complete extraction result from Agent 1."""
    # Raw content
    extracted_text: str  # Full raw text
    text_length: int
    
    # Structured data
    entities: List[ExtractedEntity] = Field(default_factory=list)
    tables: List[ExtractedTable] = Field(default_factory=list)
    metadata: Dict[str, Any] = Field(default_factory=dict)
    
    # Quality indicators
    confidence_score: float = Field(ge=0.0, le=1.0)
    ambiguities: List[Ambiguity] = Field(default_factory=list)
    corrupted_sections: List[str] = Field(default_factory=list)
    
    # OCR info (if used)
    ocr_used: bool = False
    ocr_confidence: Optional[float] = None
    
    # Source verification
    source_hash: str  # SHA-256 of input
    
    # Metadata
    extractor_version: str = "1.0.0"
    extraction_method: str  # "pypdf", "pdfplumber", "docx", "html"
    extracted_at: datetime = Field(default_factory=datetime.utcnow)


class ExtractorAgent:
    """
    Agent 1: Extractor
    
    Extracts structured data from documents with zero interpretation.
    
    Rules:
    - Extract ONLY what is explicitly present
    - Flag unclear/corrupted sections as ambiguous
    - Never clean up typos or errors
    - Preserve exact wording
    - Cite source location for every extraction
    """
    
    def __init__(self):
        self.version = "1.0.0"
    
    async def extract(
        self,
        file_content: bytes,
        filename: str,
        mime_type: Optional[str] = None
    ) -> ExtractionResult:
        """
        Extract structured data from document.
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            mime_type: MIME type (will be detected if not provided)
        
        Returns:
            ExtractionResult with all extracted data
        """
        # Calculate source hash
        source_hash = hashlib.sha256(file_content).hexdigest()
        
        # Parse filename metadata and get cleaned name
        filename_metadata = DocumentService.parse_filename_metadata(filename)
        
        # Detect MIME type if not provided
        if not mime_type:
            mime_type = magic.from_buffer(file_content, mime=True)
        
        # Route to appropriate extractor
        result = None
        if mime_type == "application/pdf":
            result = await self._extract_from_pdf(file_content, source_hash)
        elif mime_type in [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword"
        ]:
            result = await self._extract_from_docx(file_content, source_hash)
        elif mime_type == "text/plain":
            result = await self._extract_from_text(file_content, source_hash)
        elif mime_type in ["text/html", "application/xhtml+xml"]:
            result = await self._extract_from_html(file_content, source_hash)
        else:
            raise ValueError(f"Unsupported MIME type: {mime_type}")
        
        # Merge filename metadata into result
        result.metadata.update({
            'original_filename': filename,
            'filename_metadata': filename_metadata,
            'cleaned_name': filename_metadata.get('cleaned_name', filename)
        })
        
        return result
    
    async def _extract_from_pdf(
        self,
        file_content: bytes,
        source_hash: str
    ) -> ExtractionResult:
        """Extract from PDF using pdfplumber (more reliable for tables)."""
        
        extracted_text = ""
        entities = []
        tables = []
        ambiguities = []
        corrupted_sections = []
        confidence_scores = []
        
        pdf_stream = io.BytesIO(file_content)
        
        try:
            with pdfplumber.open(pdf_stream) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    # Extract text
                    page_text = page.extract_text()
                    
                    if page_text:
                        extracted_text += f"\n--- Page {page_num} ---\n{page_text}\n"
                        confidence_scores.append(1.0)
                        
                        # Extract entities from page
                        page_entities = self._extract_entities(
                            page_text,
                            location=f"page {page_num}"
                        )
                        entities.extend(page_entities)
                    else:
                        # No text extracted - might be scanned image
                        ambiguities.append(Ambiguity(
                            location=f"page {page_num}",
                            issue="No text extracted - may be scanned image requiring OCR"
                        ))
                        confidence_scores.append(0.5)
                    
                    # Extract tables
                    page_tables = page.extract_tables()
                    for table_idx, table_data in enumerate(page_tables):
                        if table_data:
                            extracted_table = self._parse_table(
                                table_data,
                                location=f"page {page_num}, table {table_idx + 1}"
                            )
                            tables.append(extracted_table)
        
        except Exception as e:
            corrupted_sections.append(f"PDF parsing error: {str(e)}")
            confidence_scores.append(0.3)
        
        # Calculate overall confidence
        overall_confidence = (
            sum(confidence_scores) / len(confidence_scores)
            if confidence_scores else 0.5
        )
        
        return ExtractionResult(
            extracted_text=extracted_text,
            text_length=len(extracted_text),
            entities=entities,
            tables=tables,
            confidence_score=overall_confidence,
            ambiguities=ambiguities,
            corrupted_sections=corrupted_sections,
            ocr_used=False,
            source_hash=source_hash,
            extraction_method="pdfplumber"
        )
    
    async def _extract_from_docx(
        self,
        file_content: bytes,
        source_hash: str
    ) -> ExtractionResult:
        """Extract from DOCX file."""
        
        extracted_text = ""
        entities = []
        tables = []
        ambiguities = []
        
        docx_stream = io.BytesIO(file_content)
        
        try:
            doc = DocxDocument(docx_stream)
            
            # Extract paragraphs
            for para_idx, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    extracted_text += para.text + "\n"
                    
                    # Extract entities
                    para_entities = self._extract_entities(
                        para.text,
                        location=f"paragraph {para_idx + 1}"
                    )
                    entities.extend(para_entities)
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                if table_data:
                    extracted_table = self._parse_table(
                        table_data,
                        location=f"table {table_idx + 1}"
                    )
                    tables.append(extracted_table)
        
        except Exception as e:
            ambiguities.append(Ambiguity(
                location="document",
                issue=f"DOCX parsing error: {str(e)}"
            ))
        
        return ExtractionResult(
            extracted_text=extracted_text,
            text_length=len(extracted_text),
            entities=entities,
            tables=tables,
            confidence_score=0.95 if not ambiguities else 0.7,
            ambiguities=ambiguities,
            source_hash=source_hash,
            extraction_method="python-docx"
        )
    
    async def _extract_from_text(
        self,
        file_content: bytes,
        source_hash: str
    ) -> ExtractionResult:
        """Extract from plain text file."""
        
        try:
            extracted_text = file_content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                extracted_text = file_content.decode('latin-1')
            except:
                return ExtractionResult(
                    extracted_text="",
                    text_length=0,
                    confidence_score=0.0,
                    corrupted_sections=["Unable to decode text file"],
                    source_hash=source_hash,
                    extraction_method="text"
                )
        
        # Extract entities from text
        entities = self._extract_entities(extracted_text, location="file")
        
        return ExtractionResult(
            extracted_text=extracted_text,
            text_length=len(extracted_text),
            entities=entities,
            confidence_score=1.0,
            source_hash=source_hash,
            extraction_method="text"
        )
    
    async def _extract_from_html(
        self,
        file_content: bytes,
        source_hash: str
    ) -> ExtractionResult:
        """Extract from HTML file (for scraped content)."""
        
        try:
            html_text = file_content.decode('utf-8')
        except UnicodeDecodeError:
            html_text = file_content.decode('latin-1')
        
        soup = BeautifulSoup(html_text, 'lxml')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Extract text
        extracted_text = soup.get_text(separator="\n", strip=True)
        
        # Extract entities
        entities = self._extract_entities(extracted_text, location="html body")
        
        # Extract tables
        tables = []
        for table_idx, table in enumerate(soup.find_all('table')):
            rows = []
            for row in table.find_all('tr'):
                cells = [cell.get_text(strip=True) for cell in row.find_all(['td', 'th'])]
                if cells:
                    rows.append(cells)
            
            if rows:
                extracted_table = self._parse_table(
                    rows,
                    location=f"html table {table_idx + 1}"
                )
                tables.append(extracted_table)
        
        return ExtractionResult(
            extracted_text=extracted_text,
            text_length=len(extracted_text),
            entities=entities,
            tables=tables,
            confidence_score=0.9,
            source_hash=source_hash,
            extraction_method="beautifulsoup"
        )
    
    def _extract_entities(
        self,
        text: str,
        location: str
    ) -> List[ExtractedEntity]:
        """
        Extract entities using regex patterns.
        
        Note: This is a simple regex-based approach. For production,
        consider using spaCy or other NER libraries.
        """
        entities = []
        
        # Money patterns ($5M, $2.5B, €100K)
        money_pattern = r'[$€£¥]\s*[\d,]+\.?\d*\s*[KMBkmb]?'
        for match in re.finditer(money_pattern, text):
            entities.append(ExtractedEntity(
                type=EntityType.MONEY,
                value=match.group(0),
                source_location=location,
                confidence=0.95,
                context=text[max(0, match.start()-20):match.end()+20]
            ))
        
        # Percentage patterns (45%, 12.5%)
        percentage_pattern = r'\d+\.?\d*\s*%'
        for match in re.finditer(percentage_pattern, text):
            entities.append(ExtractedEntity(
                type=EntityType.PERCENTAGE,
                value=match.group(0),
                source_location=location,
                confidence=0.98,
                context=text[max(0, match.start()-20):match.end()+20]
            ))
        
        # Email patterns
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        for match in re.finditer(email_pattern, text):
            entities.append(ExtractedEntity(
                type=EntityType.EMAIL,
                value=match.group(0),
                source_location=location,
                confidence=1.0,
                context=text[max(0, match.start()-20):match.end()+20]
            ))
        
        # URL patterns
        url_pattern = r'https?://[^\s<>"{}|\\^`\[\]]+'
        for match in re.finditer(url_pattern, text):
            entities.append(ExtractedEntity(
                type=EntityType.URL,
                value=match.group(0),
                source_location=location,
                confidence=1.0
            ))
        
        # Date patterns (YYYY-MM-DD, MM/DD/YYYY, etc.)
        date_patterns = [
            r'\d{4}-\d{2}-\d{2}',  # 2024-03-15
            r'\d{1,2}/\d{1,2}/\d{4}',  # 3/15/2024
            r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* \d{1,2},? \d{4}',  # March 15, 2024
        ]
        for pattern in date_patterns:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                entities.append(ExtractedEntity(
                    type=EntityType.DATE,
                    value=match.group(0),
                    source_location=location,
                    confidence=0.9,
                    context=text[max(0, match.start()-20):match.end()+20]
                ))
        
        return entities
    
    def _parse_table(
        self,
        table_data: List[List[str]],
        location: str
    ) -> ExtractedTable:
        """Parse table data into structured format."""
        
        if not table_data:
            return ExtractedTable(
                location=location,
                headers=[],
                rows=[],
                confidence=0.0,
                malformed=True
            )
        
        # Assume first row is headers
        headers = table_data[0]
        rows = table_data[1:]
        
        # Check if table is well-formed
        malformed = False
        if not all(len(row) == len(headers) for row in rows):
            malformed = True
        
        confidence = 0.95 if not malformed else 0.7
        
        return ExtractedTable(
            location=location,
            headers=headers,
            rows=rows,
            confidence=confidence,
            malformed=malformed
        )


# Singleton instance
_extractor_agent: Optional[ExtractorAgent] = None


def get_extractor_agent() -> ExtractorAgent:
    """Get or create extractor agent instance."""
    global _extractor_agent
    if _extractor_agent is None:
        _extractor_agent = ExtractorAgent()
    return _extractor_agent

